
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import ResumeInput


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(30, 30, 30)


def _add_separator(doc: Document):
    """Add a thin horizontal line via a paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    # Use a dotted line via underscored text
    run = p.add_run("─" * 70)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(180, 180, 180)


def generate_docx_bytes(
    resume: ResumeInput,
    professional_summary: str = "",
) -> bytes:
    """Generate a DOCX file from resume data. Returns raw bytes."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    
    info = resume.personal_info
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(info.name)
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(30, 30, 30)

    # Contact line
    contact_parts = []
    if info.email:
        contact_parts.append(info.email)
    if info.phone:
        contact_parts.append(info.phone)
    if info.location:
        contact_parts.append(info.location)
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(100, 100, 100)

    # Links line
    links = []
    if info.linkedin:
        links.append(f"LinkedIn: {info.linkedin}")
    if info.github:
        links.append(f"GitHub: {info.github}")
    if links:
        links_para = doc.add_paragraph()
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_run = links_para.add_run(" | ".join(links))
        links_run.font.size = Pt(9)
        links_run.font.color.rgb = RGBColor(100, 100, 100)

    
    if professional_summary:
        _add_separator(doc)
        _add_heading(doc, "Professional Summary", level=2)
        summary_para = doc.add_paragraph(professional_summary)
        summary_para.paragraph_format.space_after = Pt(4)
        for run in summary_para.runs:
            run.font.size = Pt(10)

    
    if resume.skills:
        _add_separator(doc)
        _add_heading(doc, "Technical Skills", level=2)
        skills_para = doc.add_paragraph(", ".join(resume.skills))
        for run in skills_para.runs:
            run.font.size = Pt(10)

    
    if resume.work_experience:
        _add_separator(doc)
        _add_heading(doc, "Work Experience", level=2)
        for exp in resume.work_experience:
            # Role — Company
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"{exp.role}")
            role_run.bold = True
            role_run.font.size = Pt(11)
            role_para.add_run(f"  —  {exp.company}").font.size = Pt(11)

            # Duration
            dur_para = doc.add_paragraph(exp.duration)
            dur_para.paragraph_format.space_before = Pt(0)
            dur_para.paragraph_format.space_after = Pt(2)
            for run in dur_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(120, 120, 120)
                run.italic = True

            # Bullet points
            for bullet in exp.bullets:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after = Pt(1)
                for run in bp.runs:
                    run.font.size = Pt(10)

    
    if resume.projects:
        _add_separator(doc)
        _add_heading(doc, "Projects", level=2)
        for proj in resume.projects:
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(proj.title)
            proj_run.bold = True
            proj_run.font.size = Pt(11)

            if proj.description:
                desc_para = doc.add_paragraph(proj.description)
                for run in desc_para.runs:
                    run.font.size = Pt(10)

            if proj.technologies:
                tech_para = doc.add_paragraph(f"Technologies: {', '.join(proj.technologies)}")
                tech_para.paragraph_format.space_after = Pt(2)
                for run in tech_para.runs:
                    run.font.size = Pt(9)
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)

    
    if resume.education:
        _add_separator(doc)
        _add_heading(doc, "Education", level=2)
        for edu in resume.education:
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(edu.degree)
            edu_run.bold = True
            edu_run.font.size = Pt(11)
            edu_para.add_run(f"  —  {edu.institution}").font.size = Pt(11)
            if edu.year:
                yr_para = doc.add_paragraph(edu.year)
                for run in yr_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(120, 120, 120)

    
    if resume.certifications:
        _add_separator(doc)
        _add_heading(doc, "Certifications", level=2)
        for cert in resume.certifications:
            cert_para = doc.add_paragraph()
            cert_run = cert_para.add_run(cert.title)
            cert_run.bold = True
            cert_run.font.size = Pt(10)
            extra = []
            if cert.issuer:
                extra.append(cert.issuer)
            if cert.year:
                extra.append(cert.year)
            if extra:
                cert_para.add_run(f"  —  {', '.join(extra)}").font.size = Pt(10)

    # Write to bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
